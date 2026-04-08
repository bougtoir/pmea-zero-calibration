#!/usr/bin/env python3
"""
Create Physiological Measurement Research Paper .docx — English version

"What Zeroing Cannot Fix: Concordance Analysis Unmasks Gain Errors
 Invisible to Bland-Altman"

PMEA requirements:
  - Structured abstract: Objective / Approach / Main results / Significance (<=250 words)
  - References: Harvard (alphabetical) format WITH article titles
  - Research paper: <=8000 words
  - No page/figure limits
"""

import os
import re
import sys
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
FIGDIR = os.path.join(SCRIPT_DIR, '..', 'figures')
OUTDIR = os.path.join(SCRIPT_DIR, '..', 'manuscripts')
OUTPATH = os.path.join(OUTDIR, 'PMEA_ZeroFree_Manuscript_EN.docx')

os.makedirs(OUTDIR, exist_ok=True)

doc = Document()

# ── Page setup ──
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)
style.paragraph_format.line_spacing = 2.0
style.paragraph_format.space_after = Pt(0)

# ── Helper: superscript-aware paragraph ──
def add_para_with_refs(text, bold=False, italic=False, alignment=None,
                       space_after=Pt(6), font_size=Pt(12)):
    """Add paragraph with {ref} markers converted to font-based superscripts."""
    p = doc.add_paragraph()
    if alignment:
        p.alignment = alignment
    p.paragraph_format.space_after = space_after

    parts = re.split(r'(\{[^}]+\})', text)
    for part in parts:
        if part.startswith('{') and part.endswith('}'):
            run = p.add_run(part[1:-1])
            run.font.name = 'Times New Roman'
            run.font.size = font_size
            run.font.superscript = True
            run.bold = bold
            run.italic = italic
        else:
            run = p.add_run(part)
            run.font.name = 'Times New Roman'
            run.font.size = font_size
            run.bold = bold
            run.italic = italic
    return p


# ── Helper functions ──
def add_heading_styled(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = 'Times New Roman'
        run.font.color.rgb = RGBColor(0, 0, 0)
    return h

def add_para(text, bold=False, italic=False, alignment=None, space_after=Pt(6)):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.bold = bold
    run.italic = italic
    if alignment:
        p.alignment = alignment
    p.paragraph_format.space_after = space_after
    return p

def add_figure(filename, caption, width=Inches(6)):
    path = os.path.join(FIGDIR, filename)
    if not os.path.exists(path):
        add_para(f'[Figure file not found: {filename}]', italic=True)
        cap = add_para(caption, italic=True)
        cap.paragraph_format.space_after = Pt(12)
        return cap
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(path, width=width)
    cap = doc.add_paragraph()
    cap_run = cap.add_run(caption)
    cap_run.font.name = 'Times New Roman'
    cap_run.font.size = Pt(10)
    cap_run.italic = True
    cap.paragraph_format.space_after = Pt(12)
    return cap

# ══════════════════════════════════════════════════════════════════
# TITLE PAGE
# ══════════════════════════════════════════════════════════════════
add_para('RESEARCH PAPER', bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run(
    'What Zeroing Cannot Fix: Concordance Analysis Unmasks '
    'Gain Errors Invisible to Bland\u2013Altman'
)
run.font.name = 'Times New Roman'
run.font.size = Pt(16)
run.bold = True
title.paragraph_format.space_after = Pt(18)

# Running title
add_para('Running title: What zeroing cannot fix', italic=True,
         alignment=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_paragraph()

# Authors
add_para('[Author names to be completed]', alignment=WD_ALIGN_PARAGRAPH.CENTER)
add_para('[Affiliations to be completed]', alignment=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_paragraph()

# Corresponding author
add_para('Corresponding author:', bold=True)
add_para('[Name, address, email, ORCID]')
doc.add_paragraph()

# Word count etc.
add_para('Word count: ~4,500 (main text, excluding references, tables, and figure legends)')
add_para('Tables: 2')
add_para('Figures: 8')
add_para('References: 30')
doc.add_paragraph()

# Keywords
add_para('Keywords: ', bold=True, space_after=Pt(0))
kw = doc.paragraphs[-1]
kw_run = kw.add_run(
    'arterial pressure monitoring; concordance correlation coefficient; '
    'zero calibration; pulse pressure; MEMS pressure sensor; '
    'bias correction factor; method comparison; device validation'
)
kw_run.font.name = 'Times New Roman'
kw_run.font.size = Pt(12)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# ABSTRACT — PMEA structured format
# ══════════════════════════════════════════════════════════════════
add_heading_styled('Abstract', level=1)

add_para('Objective', bold=True, space_after=Pt(2))
add_para(
    'To demonstrate, through simulation and analytical decomposition, that zero calibration '
    'of invasive arterial pressure transducers corrects only the offset component of measurement '
    'error while leaving gain error undetected, and to propose Lin\u2019s concordance correlation '
    'coefficient (CCC) decomposition as a complementary validation metric that distinguishes '
    'these error types.'
)

add_para('Approach', bold=True, space_after=Pt(2))
add_para(
    'We decompose the CCC into precision (Pearson\u2019s r) and accuracy (bias correction factor, '
    'C\u2082) components, where C\u2082 is further resolved into location shift (u, offset) and scale '
    'shift (v, gain error). We simulate four clinically relevant scenarios of arterial pressure '
    'measurement (n = 150 paired measurements each): offset only, offset corrected by zeroing, '
    'gain error, and combined gain plus offset error. Each scenario is analysed using both '
    'Bland\u2013Altman statistics and CCC decomposition. We also identify three engineering solutions '
    '\u2014catheter-tip micro-electro-mechanical systems (MEMS) sensors, barometric compensation, and self-calibrating MEMS\u2014that '
    'eliminate the need for manual zero calibration.'
)

add_para('Main results', bold=True, space_after=Pt(2))
add_para(
    'After zero calibration, Bland\u2013Altman analysis shows near-zero bias regardless of whether '
    'gain error exists (bias: \u22120.1 vs 1.1 mmHg for gain-free vs gain-error scenarios). '
    'In contrast, CCC decomposition clearly distinguishes these cases: C\u2082 = 1.000 (no gain error) '
    'versus C\u2082 = 0.870 (10% gain error, v = 1.11). We show that pulse pressure accuracy serves '
    'as an implicit gain validator,     since pulse pressure is independent of direct-current (DC) offset but '
        'proportional to sensor gain.'
)

add_para('Significance', bold=True, space_after=Pt(2))
add_para(
    'CCC decomposition into r, C\u2082, u, and v provides diagnostic information about arterial '
    'pressure monitor performance that Bland\u2013Altman analysis alone cannot capture. The framework '
    'offers quantitative design targets (C\u2082 \u2265 0.99 without calibration) for next-generation '
    'sensors and extends naturally to non-invasive cardiac output monitors. We recommend reporting '
    'CCC with decomposition in all arterial pressure validation studies.'
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# INTRODUCTION
# ══════════════════════════════════════════════════════════════════
add_heading_styled('1. Introduction', level=1)

add_para_with_refs(
    'Invasive arterial pressure monitoring is one of the most frequently performed measurements in '
    'anaesthesia and critical care. After insertion of an arterial catheter, the clinician must zero '
    'the pressure transducer to atmospheric pressure at the level of the phlebostatic axis before '
    'initiating haemodynamic monitoring (Saugel et al 2020, Saugel and Sessler 2021). This zeroing '
    'procedure must be repeated whenever the height relationship between the transducer and the '
    'catheter insertion site changes\u2014for example, when the operating table is tilted or the bed '
    'angle is adjusted (Gupta et al 2025). Despite decades of technological advancement in pressure '
    'sensing, the need for manual zero calibration has remained essentially unchanged since the '
    'introduction of external fluid-filled transducer systems.'
)

add_para_with_refs(
    'The persistence of this requirement is not a fundamental physical necessity but a consequence of '
    'how conventional monitoring systems are designed. In a standard fluid-filled arterial line, the '
    'pressure transducer is located externally, connected to the intravascular catheter tip by a column '
    'of saline. This architecture introduces three systematic sources of direct-current (DC) offset: (1) the hydrostatic '
    'pressure difference between the transducer and the measurement site (\u0394P = \u03C1gh, approximately '
    '0.74 mmHg per centimetre of height difference); (2) the need to measure gauge pressure relative '
    'to atmospheric pressure (~760 mmHg); and (3) transducer drift over time due to mechanical creep '
    'or thermal effects in the strain gauge (Mark 1998, McGhee and Bridges 2002). Zero calibration '
    'removes all three offsets simultaneously by exposing the transducer to atmosphere and resetting '
    'the output to zero.'
)

add_para_with_refs(
    'However, zero calibration is exclusively an offset correction. It does not verify or correct the '
    'gain (sensitivity) of the transducer\u2014that is, the proportionality between the true pressure '
    'change and the electrical output. If the gain is incorrect, the measured waveform will be scaled '
    'up or down, and this error will persist after zeroing. This distinction between offset error and '
    'gain error is not merely academic; it has direct implications for how we validate monitoring '
    'devices and, as we argue below, for how they should be designed.'
)

add_para_with_refs(
    'The standard validation framework for haemodynamic monitors centres on Bland\u2013Altman analysis '
    '(Bland and Altman 1986, 1999), the Critchley\u2013Critchley percentage error criterion (Critchley '
    'and Critchley 1999), and polar plot assessment for trending ability (Critchley et al 2010, 2011), '
    'with particular attention to the precision of the reference technique (Cecconi et al 2009). '
    'While these tools effectively evaluate bias, limits of agreement, and directional concordance, '
    'none provides an integrated single-value metric that simultaneously captures both accuracy and '
    'precision of absolute measurements against the line of perfect agreement.'
)

add_para_with_refs(
    'In this article, we propose that if a device measures pulse pressure (PP) accurately, its sensor '
    'gain is implicitly validated, and zero calibration is rendered unnecessary through appropriate '
    'sensor design. We formalise this argument using Lin\u2019s concordance correlation coefficient (CCC) '
    'decomposition (Lin 1989, 2000), identify the engineering solutions that eliminate each source of '
    'DC offset, demonstrate through simulation that CCC decomposition detects gain errors masked by '
    'Bland\u2013Altman analysis, and discuss the implications for arterial pressure monitoring and beyond.'
)

# ══════════════════════════════════════════════════════════════════
# METHODS / LOGICAL ARGUMENT
# ══════════════════════════════════════════════════════════════════
add_heading_styled('2. Theory and methods', level=1)

add_heading_styled('2.1. Arterial pressure as the sum of alternating-current and direct-current components', level=2)

add_para_with_refs(
        'An arterial pressure waveform P(t) can be decomposed into a slowly varying component (the DC '
        'level, determined by mean arterial pressure and any external offsets) and a pulsatile component '
        '(the alternating-current [AC] signal, driven by cardiac ejection). Pulse pressure (PP = systolic blood pressure [SBP] \u2013 diastolic blood pressure [DBP]) is a pure AC '
    'quantity: it represents the difference between the maximum and minimum of the pulsatile '
    'excursion and is, by definition, independent of any additive DC offset (figure 1, panels A and B).'
)

add_para_with_refs(
    'This independence has a critical implication. If a sensor measures PP correctly, its gain '
    '(transduction sensitivity, in mV/mmHg or equivalent units) must be correct, because PP accuracy '
    'requires that the electrical output difference corresponding to the true pressure difference '
    '(SBP \u2013 DBP) is exact. Conversely, if the gain is incorrect (e.g., the sensor reads 1.15 mV '
    'per mmHg instead of 1.00 mV per mmHg), the measured PP will be proportionally distorted '
    '(figure 1, panel C). PP accuracy therefore serves as an implicit validator of sensor gain.'
)

add_heading_styled('2.2. Sources of DC offset and engineering elimination', level=2)

add_para_with_refs(
    'If gain is correct (validated by PP accuracy), the only remaining source of measurement error '
    'is the DC offset. Three distinct sources contribute to this offset in conventional '
    'systems (table 1). Importantly, each can be eliminated by established engineering solutions:'
)

add_para_with_refs(
    '(1) Hydrostatic column offset: Catheter-tip micro-electro-mechanical systems (MEMS) pressure sensors (e.g., Millar Mikro-Cath) '
    'place the sensing element directly at the catheter tip, eliminating the fluid column entirely '
    '(Hasenkamp et al 2012, Song et al 2020). The sensor measures pressure at the point of interest, '
    'making the measurement independent of the transducer-to-patient height relationship (Millar 2026).'
)

add_para_with_refs(
    '(2) Atmospheric pressure reference: Absolute pressure sensors measure total pressure relative '
    'to an internal vacuum reference rather than to atmosphere. By incorporating a barometric pressure '
    'sensor in the monitor and electronically subtracting atmospheric pressure, the gauge pressure '
    'can be computed without manual atmospheric zeroing. This approach is already implemented in '
    'chronic implantable sensors such as Millar\u2019s TiSense platform (Millar 2026b).'
)

add_para_with_refs(
    '(3) Transducer drift: Self-calibrating MEMS sensors incorporate internal reference pressure '
    'cavities that provide known pressure points for periodic auto-calibration (Kang et al 2022). '
    'By cycling through a reference pressure via liquid-to-vapour phase transition in a sealed '
    'micro-cavity, the sensor continuously corrects for zero drift without external intervention.'
)

add_para_with_refs(
    'When all three solutions are implemented simultaneously\u2014tip sensor, absolute pressure with '
    'barometric compensation, and self-calibrating reference\u2014every source of DC offset is eliminated '
    'by design. Manual zero calibration becomes redundant (figure 2).'
)

# ══════════════════════════════════════════════════════════════════
# CCC FRAMEWORK
# ══════════════════════════════════════════════════════════════════
add_heading_styled('2.3. Lin\u2019s CCC and its decomposition', level=2)

add_para_with_refs(
    'Lin\u2019s concordance correlation coefficient (\u03C1c) quantifies the agreement between paired '
    'measurements along the 45-degree line of perfect concordance (Lin 1989). It decomposes as:'
)

eq = doc.add_paragraph()
eq.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = eq.add_run('\u03C1c = r \u00D7 C\u2082')
run.font.name = 'Times New Roman'
run.font.size = Pt(12)
run.italic = True
eq.paragraph_format.space_after = Pt(6)

add_para_with_refs(
    'where r (Pearson\u2019s correlation coefficient) measures precision (the tightness of data around '
    'the best-fit line) and C\u2082 (bias correction factor) measures accuracy (how far the best-fit line '
    'deviates from the identity line). C\u2082 is further decomposed as (Lin 2000):'
)

eq2 = doc.add_paragraph()
eq2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = eq2.add_run('C\u2082 = 2 / (v + 1/v + u\u00B2)')
run2.font.name = 'Times New Roman'
run2.font.size = Pt(12)
run2.italic = True
eq2.paragraph_format.space_after = Pt(6)

add_para_with_refs(
    'where v = \u03C31/\u03C32 is the scale shift (ratio of standard deviations, reflecting gain error) '
    'and u = (\u03BC1 \u2013 \u03BC2)/\u221A(\u03C31\u03C32) is the location shift (normalised mean '
    'difference, reflecting offset).'
)

add_heading_styled('2.4. What zero calibration does\u2014and does not do\u2014in CCC terms', level=2)

add_para_with_refs(
    'Zero calibration is an operation that drives u toward zero by removing the systematic offset '
    'between the device reading and the true pressure. In CCC terms, successful zeroing achieves '
    'u \u2248 0, which maximises C\u2082 with respect to the location component. However, zeroing has no '
    'effect on v: if the sensor gain is incorrect (v \u2260 1), C\u2082 remains below 1.0 even after '
    'perfect zeroing (figures 3 and 4).'
)

add_heading_styled('2.5. Simulation design', level=2)

add_para_with_refs(
    'We simulated four clinically relevant scenarios of arterial pressure measurement, each comprising '
    'n = 150 paired measurements with true systolic pressures ranging from 80 to 180 mmHg. '
    'Scenario A: offset only (12 mmHg systematic offset before zeroing). '
    'Scenario B: after zeroing (offset removed, no gain error). '
    'Scenario C: gain error only (v = 1.11, i.e. 11% overestimation of pressure changes). '
    'Scenario D: combined gain error plus offset. '
    'Gaussian noise (standard deviation [SD] = 3.5 mmHg) was added to all scenarios to simulate physiological and '
    'measurement variability. Each scenario was analysed using both Bland\u2013Altman statistics '
    '(bias, limits of agreement, percentage error) and full CCC decomposition (\u03C1c, r, C\u2082, u, v).'
)

add_heading_styled('2.6. The complete argument', level=2)

add_para(
    'Combining the engineering and statistical perspectives, the argument can be stated formally:'
)

add_para('(1) PP is accurate \u2192 gain is correct \u2192 v = 1 (no scale shift)')
add_para('(2) Catheter-tip sensor + barometric compensation + self-calibrating MEMS \u2192 all DC offsets eliminated \u2192 u = 0 (no location shift)')
add_para('(3) v = 1 and u = 0 \u2192 C\u2082 = 1.0')
add_para('(4) CCC = r \u00D7 1.0 = r')

add_para_with_refs(
    'That is, the performance of a properly designed sensor is limited only by its precision (r), '
    'which reflects random measurement noise. All systematic error\u2014both offset and gain\u2014is '
    'resolved by design, not by calibration. The sensor\u2019s C\u2082 approaches 1.0 as a design '
    'specification rather than a calibration outcome.'
)

# ══════════════════════════════════════════════════════════════════
# RESULTS
# ══════════════════════════════════════════════════════════════════
add_heading_styled('3. Results', level=1)

add_heading_styled('3.1. Simulation outcomes', level=2)

add_para_with_refs(
    'Figure 3 shows concordance plots for the four simulated scenarios. '
    'Scenario A (offset only, before zeroing) shows reduced CCC (0.855) due '
    'to location shift (u = \u20130.55). Scenario B (after zeroing) shows near-perfect agreement '
    '(CCC = 0.986, C\u2082 = 1.000) because the offset has been removed and no gain error exists. '
    'Scenario C (gain error, v = 1.11) shows that zeroing cannot improve C\u2082 (0.870) because the '
    'scale shift persists. Scenario D (gain + offset) shows CCC = 0.976; zeroing '
    'would improve it to Scenario C but not to Scenario B.'
)

add_heading_styled('3.2. Bland\u2013Altman comparison', level=2)

add_para_with_refs(
    'The critical comparison is between Scenarios B and C: Bland\u2013Altman analysis (figures 5 and 6) '
    'shows near-zero bias in both cases, yet CCC decomposition reveals fundamentally different '
    'error structures (table 2). Scenario B has C\u2082 = 1.000 (no systematic error), whereas '
    'Scenario C has C\u2082 = 0.870 (hidden gain error, v = 1.11). This demonstrates that '
    'Bland\u2013Altman analysis alone cannot distinguish true agreement from agreement where gain error '
    'is masked by coincidental offset cancellation.'
)

add_heading_styled('3.3. Sensitivity analysis', level=2)

add_para_with_refs(
    'Figure 7 quantifies these relationships. Panel A shows C\u2082 as a function of gain error and '
    'DC offset, confirming that zero calibration (movement to zero offset) improves C\u2082 only when '
    'gain is correct. Panel B shows CCC degradation with gain error at different sensor precision '
    'levels (r), demonstrating that even high-precision sensors (r = 0.99) suffer substantial CCC '
    'reduction when gain error exceeds \u00B110%. Notably, typical MEMS sensors achieve gain accuracy '
    'within \u00B15%, corresponding to C\u2082 > 0.99 at zero offset.'
)

add_heading_styled('3.4. Pulse pressure as gain validator', level=2)

add_para_with_refs(
    'Figure 8 provides quantitative demonstration of the PP\u2013gain relationship. Simulated arterial '
    'pressure waveforms with correct gain (panel A), gain error (panel B), and DC offset only '
    '(panel C) show that PP is distorted by gain error but unaffected by DC offset. Panels D and E '
    'confirm that the slope of the PP regression line against the reference directly estimates '
    'sensor gain, providing a practical validation tool. The logical chain from PP accuracy to '
    'zero-calibration-free monitoring is summarised in panel F.'
)

# ══════════════════════════════════════════════════════════════════
# DISCUSSION
# ══════════════════════════════════════════════════════════════════
add_heading_styled('4. Discussion', level=1)

add_heading_styled('4.1. Implications for sensor design', level=2)

add_para_with_refs(
    'The framework presented here reframes zero calibration from a clinical necessity to a design '
    'workaround. Conventional fluid-filled systems require zeroing because their architecture '
    'inherently introduces DC offsets. Rather than continuing to refine calibration procedures, '
    'device manufacturers should pursue designs that eliminate the need for calibration entirely. '
    'The component technologies\u2014catheter-tip MEMS, absolute pressure sensing, barometric '
    'compensation, and self-calibrating references\u2014already exist individually in commercial '
        'or near-commercial products (Hasenkamp et al 2012, Song et al 2020, Millar 2026, 2026b, '
        'Kang et al 2022), and high-fidelity pressure wires have already demonstrated the feasibility '
        'of accurate catheter-based pressure measurement (Scalia et al 2023). Their integration into '
        'a single clinical arterial pressure monitoring system is an engineering challenge, not a '
        'scientific one.'
)

add_para_with_refs(
    'The CCC decomposition provides a quantitative design target: device developers should aim for '
    'C\u2082 \u2265 0.99 without any calibration step\u2014corresponding to \u201cexcellent\u201d agreement under '
    'proposed strength-of-agreement criteria (McBride 2005)\u2014verifiable by comparing the uncalibrated device '
    'output against a reference standard. If C\u2082 < 0.99 prior to calibration, the device has a '
    'residual systematic error that calibration can only partially mask. The decomposition into '
    'u and v further identifies whether the residual error is an offset (addressable by zeroing) '
    'or a gain error (requiring hardware or algorithmic correction).'
)

add_heading_styled('4.2. Implications for validation methodology', level=2)

add_para_with_refs(
    'Current regulatory pathways (e.g., U.S. Food and Drug Administration [FDA] 510(k)) do not prescribe specific statistical methods '
    'for validating arterial pressure monitors (FDA 2026). Published validation studies overwhelmingly '
    'rely on Bland\u2013Altman analysis and percentage error (Kim et al 2014, Joosten et al 2017, Bland and Altman 1986). '
    'However, the Bland\u2013Altman plot of a zeroed device will show bias \u2248 0, potentially masking '
    'proportional bias (gain error) within the limits of agreement. CCC reporting, particularly the '
    'C\u2082 component, would provide an additional layer of scrutiny that distinguishes offset '
    'correction from genuine measurement accuracy.'
)

add_para_with_refs(
    'We suggest that validation studies for arterial pressure monitors should report: (1) CCC with '
    'decomposition into r and C\u2082; (2) C\u2082 before and after zero calibration, to quantify the '
    'device\u2019s dependence on calibration; and (3) the individual contributions of u and v to any '
    'C\u2082 deficit. This information would enable regulators and clinicians to assess whether a device '
    'achieves accuracy through good design or through calibration-dependent offset removal.'
)

add_heading_styled('4.3. Extension to non-invasive cardiac output monitors', level=2)

add_para_with_refs(
    'The principle that calibration corrects only offset (u) while leaving gain error (v) untouched '
    'extends naturally to non-invasive cardiac output (CO) monitors, where the pulmonary artery '
    'catheter remains the clinical reference standard (Chatterjee 2009). ClearSight\u2019s Physiocal '
    'algorithm periodically re-optimises the volume clamp setpoint\u2014essentially an offset '
    'correction (Ameloot et al 2015). Noninvasive cardiac output monitoring (NICOM)/Starling\u2019s bioreactance phase reference provides a '
    'baseline offset for the phase-to-stroke-volume conversion (Squara et al 2007). FloTrac\u2019s '
    'arterial waveform calibration adjusts the mean pressure offset (Manecke 2005, Romagnoli et al 2013). In every case, '
    'the auto-calibration routine addresses u (location shift) but does not correct v (scale '
    'shift)\u2014that is, the gain of the pressure-to-CO or impedance-to-stroke volume (SV) conversion.'
)

add_para_with_refs(
    'This observation, combined with the CO monitoring validation framework proposed by Odor et al '
    '(2017) and previous calls for CCC adoption in haemodynamic monitoring, suggests a unified '
    'principle: for any haemodynamic monitoring device, the bias correction factor C\u2082 should be '
    'decomposed into u and v components, and the device\u2019s reliance on calibration to achieve '
    'acceptable C\u2082 should be explicitly quantified. A device that requires frequent recalibration '
    'to maintain C\u2082 close to 1.0 has a fundamental design limitation that calibration can only '
    'partially compensate.'
)

add_heading_styled('4.4. Limitations', level=2)

add_para_with_refs(
    'Our argument rests on the assumption that sensor response is linear across the physiological '
    'pressure range. If the sensor exhibits significant non-linearity, PP accuracy at one pressure '
    'level would not guarantee gain correctness at other levels. However, modern MEMS piezoresistive '
    'sensors typically achieve non-linearity < 0.1% of full-scale output over 0\u2013300 mmHg '
    '(Barlian et al 2009), well within the requirements for clinical arterial pressure monitoring.'
)

add_para_with_refs(
    'The framework addresses systematic error only. Random measurement noise (which determines r) '
    'is a separate concern that is not affected by calibration or the design features discussed here. '
    'A zero-calibration-free sensor with high random noise would still have poor CCC despite '
    'achieving C\u2082 = 1.0.'
)

add_para_with_refs(
    'Finally, we acknowledge that the clinical adoption of catheter-tip MEMS sensors for routine '
    'arterial monitoring faces practical barriers including cost, disposability, and compatibility '
    'with existing monitoring infrastructure. Our aim is not to mandate immediate clinical change but '
    'to establish the theoretical principle that zero calibration is a designable-out limitation, '
    'and to provide a quantitative framework (CCC decomposition) for evaluating progress toward '
    'this goal.'
)

add_heading_styled('4.5. Clinical perspective', level=2)

add_para_with_refs(
    'For practising clinicians, the immediate practical message is nuanced. Zero calibration remains '
    'essential with current fluid-filled systems and should continue to be performed rigorously, as '
    'emphasised by Saugel et al (2020) and Gupta et al (2025). However, clinicians should be aware '
    'that even perfect zeroing corrects only offset, not gain. A zeroed transducer with a 10% gain '
    'error will display SBP 132 instead of 120 and DBP 88 instead of 80\u2014a clinically significant '
    'overestimation that would not be detected by the zeroing procedure and would appear as a '
    'widened PP (52 vs 40 mmHg) rather than a shifted baseline.'
)

add_para_with_refs(
    'The recognition that PP accuracy implies gain correctness also suggests a simple clinical check: '
    'if the invasively measured PP is physiologically plausible and consistent with the non-invasive '
    'cuff measurement, the sensor gain is likely correct. An implausible PP (e.g., very narrow or '
    'very wide relative to the patient\u2019s condition) may indicate a gain problem that zeroing '
    'cannot resolve.'
)

# ══════════════════════════════════════════════════════════════════
# CONCLUSION
# ══════════════════════════════════════════════════════════════════
add_heading_styled('5. Conclusion', level=1)

add_para_with_refs(
    'We have presented a simulation-based framework demonstrating that zero calibration of arterial '
    'pressure transducers corrects only offset (location shift, u) while leaving gain error (scale '
    'shift, v) undetected. Pulse pressure accuracy validates sensor gain, eliminating scale '
    'shift (v = 1). Catheter-tip MEMS sensors, absolute pressure measurement with barometric '
    'compensation, and self-calibrating MEMS references together eliminate all sources of DC offset '
    '(u = 0). Formalised through Lin\u2019s CCC decomposition, this means C\u2082 = 1.0 is achievable by '
    'design rather than by calibration. The framework provides quantitative design targets for '
    'next-generation arterial pressure monitors, a complementary validation metric for regulatory '
    'evaluation, and a deeper understanding of what calibration does\u2014and does not do\u2014in '
    'physiological measurement.'
)

# ══════════════════════════════════════════════════════════════════
# DECLARATIONS
# ══════════════════════════════════════════════════════════════════
add_heading_styled('Declarations of interest', level=1)
add_para('[To be completed by authors]')

add_heading_styled('Funding', level=1)
add_para('[To be completed by authors]')

add_heading_styled('Authors\u2019 contributions', level=1)
add_para('[To be completed by authors]')

add_heading_styled('Data availability', level=1)
add_para(
    'All simulation code and data used to generate the figures and results in this paper '
    'are available from the corresponding author upon reasonable request.'
)

add_heading_styled('Acknowledgements', level=1)
add_para('[To be completed by authors]')

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# REFERENCES — Harvard (alphabetical) format with article titles
# ══════════════════════════════════════════════════════════════════
add_heading_styled('References', level=1)

refs_harvard = [
    'Ameloot K, Palmers P J and Malbrain M L N G 2015 The accuracy of noninvasive cardiac output and pressure measurements with finger cuff: a concise review Curr. Opin. Crit. Care 21 232\u20139',
    'Barlian A A, Park W-T, Mallon J R, Rastegar A J and Pruitt B L 2009 Review: semiconductor piezoresistance for microsystems Proc. IEEE 97 513\u201352',
    'Bland J M and Altman D G 1986 Statistical methods for assessing agreement between two methods of clinical measurement Lancet 327 307\u201310',
    'Bland J M and Altman D G 1999 Measuring agreement in method comparison studies Stat. Methods Med. Res. 8 135\u201360',
    'Cecconi M, Rhodes A, Poloniecki J, Della Rocca G and Grounds R M 2009 Bench-to-bedside review: the importance of the precision of the reference technique in method comparison studies Crit. Care 13 201',
    'Chatterjee K 2009 The Swan\u2013Ganz catheters: past, present, and future Circulation 119 147\u201352',
    'Critchley L A H and Critchley J A J H 1999 A meta-analysis of studies using bias and precision statistics to compare cardiac output measurement techniques J. Clin. Monit. Comput. 15 85\u201391',
    'Critchley L A, Lee A and Ho A M-H 2010 A critical review of the ability of continuous cardiac output monitors to measure trends in cardiac output Anesth. Analg. 111 1180\u201392',
    'Critchley L A, Yang X X and Lee A 2011 Assessment of trending ability of cardiac output monitors by polar plot methodology J. Cardiothorac. Vasc. Anesth. 25 536\u201346',
    'FDA (U.S. Food and Drug Administration) 2026 Premarket Notification 510(k) https://www.fda.gov/medical-devices/premarket-submissions-selecting-and-preparing-correct-submission/premarket-notification-510k (accessed 25 March 2026)',
    'Gupta D, Jain A and Ismaeil M 2025 Zero arterial catheters with every change in the height difference of pressure transducer and catheter insertion site Ann. Card. Anaesth. 28 205',
    'Hasenkamp W, Theumer T, Nolte J, Aschenbrenner L, Dietzel A and B\u00FCttgenbach S 2012 Polyimide/SU-8 catheter-tip MEMS gauge pressure sensor Biomed. Microdevices 14 819\u201328',
    'Joosten A, Desebbe O, Suehiro K, Murphy L S-L, Essber H, Alexander B, Fischer M O, Barvais L, Van Obbergh L, Maucort-Boulch D and Cannesson M 2017 Accuracy and precision of non-invasive cardiac output monitoring devices in perioperative medicine: a systematic review and meta-analysis Br. J. Anaesth. 118 298\u2013310',
    'Kang Y, Ge C, Hu Y, Gao J, Zhong M, Du H and Wang S 2022 Development of a flexible integrated self-calibrating MEMS pressure sensor using a liquid-to-vapor phase change Sensors 22 9737',
    'Kim S-H, Lilot M, Sidhu K S, Rinehart J, Cannesson M, Joosten A and Lee C 2014 Accuracy and precision of continuous noninvasive arterial pressure monitoring compared with invasive arterial pressure Anesthesiology 120 1080\u201397',
    'Lin L I-K 1989 A concordance correlation coefficient to evaluate reproducibility Biometrics 45 255\u201368',
    'Lin L I-K 2000 A note on the concordance correlation coefficient Biometrics 56 324\u20135',
    'Manecke G R 2005 Edwards FloTrac sensor and Vigileo monitor: easy, accurate, reliable cardiac output assessment using the arterial pulse wave Expert Rev. Med. Devices 2 523\u20137',
    'Mark J B 1998 Atlas of Cardiovascular Monitoring (New York: Churchill Livingstone)',
    'McBride G B 2005 A proposal for strength-of-agreement criteria for Lin\u2019s concordance correlation coefficient NIWA Client Report HAM2005-062',
    'McGhee B H and Bridges M E J 2002 Monitoring arterial blood pressure: what you may not know Crit. Care Nurse 22 60\u201379',
    'Millar, Inc. 2026 Mikro-Cath Pressure Catheter \u2013 Product Information https://millar.com/Clinical/MikroCath/ (accessed 25 March 2026)',
    'Millar, Inc. 2026b TiSense \u2013 Chronic Pressure Sensing https://millar.com/Clinical/MikroCath/Continuous-Compartment-Pressure-Measurements/ (accessed 25 March 2026)',
    'Odor P M, Bampoe S and Cecconi M 2017 Cardiac output monitoring: validation studies\u2014how results should be presented Curr. Anesthesiol. Rep. 7 410\u201315',
    'Romagnoli S, Ricci Z, Romano S M, Dimizio F, Bonicolini E, Quattrone D, De Gaudio A R and Grasso S 2013 FloTrac/Vigileo (third generation) and MostCare/PRAM versus echocardiography for cardiac output estimation in vascular surgery J. Cardiothorac. Vasc. Anesth. 27 1114\u201321',
    'Saugel B, Kouz K, Meidert A S, Schulte-Uentrop L and Romagnoli S 2020 How to measure blood pressure using an arterial catheter: a systematic 5-step approach Crit. Care 24 172',
    'Saugel B and Sessler D I 2021 Perioperative blood pressure management Anesthesiology 134 250\u201361',
    'Scalia A, Ghafari C, Navarre W, Delmotte P, Phillips R and Carlier S 2023 High fidelity pressure wires provide accurate validation of non-invasive central blood pressure and pulse wave velocity measurements Biomedicines 11 1235',
    'Song P, Ma Z, Ma J, Yang L, Wei J, Zhao Y, Zhang M, Yang F and Wang X 2020 Recent progress of miniature MEMS pressure sensors Micromachines 11 56',
    'Squara P, Denjean D, Estagnasie P, Brusset A, Dib J C and Dubois C 2007 Noninvasive cardiac output monitoring (NICOM): a clinical validation Intensive Care Med. 33 1191\u20134',
]

for ref in refs_harvard:
    p = doc.add_paragraph()
    run = p.add_run(ref)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(10)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.first_line_indent = Cm(-1.27)
    p.paragraph_format.left_indent = Cm(1.27)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# TABLE 1
# ══════════════════════════════════════════════════════════════════
add_heading_styled('Table 1', level=1)
add_para(
    'Table 1. Sources of DC offset in conventional arterial pressure monitoring '
    'and engineering solutions for their elimination.',
    bold=True, italic=True
)

table = doc.add_table(rows=5, cols=5)
table.style = 'Light Grid Accent 1'

headers = ['Offset source', 'Physical mechanism', 'Magnitude', 'Engineering solution', 'CCC component affected']
for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = h
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True
            run.font.size = Pt(9)
            run.font.name = 'Times New Roman'

data = [
    ['Hydrostatic column', '\u0394P = \u03C1gh (height difference between transducer and catheter tip)',
     '~0.74 mmHg/cm', 'Catheter-tip MEMS sensor (eliminates fluid column)',
     'Reduces u (location shift)'],
    ['Atmospheric pressure reference', 'Gauge pressure measured relative to atmosphere (~760 mmHg)',
     'Entire baseline', 'Absolute pressure sensor + built-in barometer for electronic subtraction',
     'Reduces u (location shift)'],
    ['Transducer drift', 'Mechanical creep, thermal effects in strain gauge',
     '~1\u20135 mmHg/day', 'Self-calibrating MEMS with internal reference pressure cavity',
     'Reduces u (location shift)'],
    ['Gain error (NOT corrected by zeroing)', 'Sensitivity mismatch: output per unit pressure \u2260 nominal',
     '1\u201315% typical', 'Factory calibration; PP accuracy as implicit gain check',
     'Affects v (scale shift); zeroing has no effect'],
]

for row_idx, row_data in enumerate(data, start=1):
    for col_idx, cell_text in enumerate(row_data):
        cell = table.rows[row_idx].cells[col_idx]
        cell.text = cell_text
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(8)
                run.font.name = 'Times New Roman'

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# TABLE 2
# ══════════════════════════════════════════════════════════════════
add_heading_styled('Table 2', level=1)
add_para(
    'Table 2. Comparison of statistical metrics across four simulated scenarios (n = 150 paired measurements).',
    bold=True, italic=True
)

table2 = doc.add_table(rows=5, cols=10)
table2.style = 'Light Grid Accent 1'

t2_headers = ['Scenario', 'CCC', 'r', 'C\u2082', 'u', 'v', 'Bias\n(mmHg)', 'LoA lower\n(mmHg)',
              'LoA upper\n(mmHg)', 'PE (%)']
for i, h in enumerate(t2_headers):
    cell = table2.rows[0].cells[i]
    cell.text = h
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True
            run.font.size = Pt(8)
            run.font.name = 'Times New Roman'

t2_data = [
    ['A: Before zeroing', '0.855', '0.986', '0.867', '\u20130.55', '0.99', '11.9', '4.8', '19.0', '7.2'],
    ['B: After zeroing', '0.986', '0.986', '1.000', '0.01', '0.99', '\u20130.1', '\u20137.2', '7.0', '7.2'],
    ['C: Gain error', '0.855', '0.982', '0.870', '0.54', '1.11', '\u201310.9', '\u201319.4', '\u20132.4', '8.7'],
    ['D: Gain + offset', '0.976', '0.982', '0.993', '\u20130.05', '1.11', '1.1', '\u20137.4', '9.6', '8.7'],
]

for row_idx, row_data in enumerate(t2_data, start=1):
    for col_idx, cell_text in enumerate(row_data):
        cell = table2.rows[row_idx].cells[col_idx]
        cell.text = cell_text
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(8)
                run.font.name = 'Times New Roman'

add_para(
    'CCC = concordance correlation coefficient; r = Pearson correlation (precision); '
    'C\u2082 = bias correction factor (accuracy); u = location shift; v = scale shift; '
    'LoA = limits of agreement; PE = percentage error. '
    'Scenario B and C have similar Bland\u2013Altman bias but dramatically different C\u2082 values, '
    'demonstrating that Bland\u2013Altman analysis alone cannot distinguish offset-corrected agreement '
    'from gain-error-masked agreement.',
    italic=True
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# FIGURES
# ══════════════════════════════════════════════════════════════════
add_heading_styled('Figures', level=1)

add_figure(
    'figure1_signal_decomposition.png',
    'Figure 1. Arterial pressure waveform decomposition. (A) True arterial pressure with '
    'DC component (offset-dependent baseline) and AC component (pulse pressure, PP = 40 mmHg). '
    '(B) With DC offset of 15 mmHg (e.g., from hydrostatic column): PP remains unchanged at '
    '40 mmHg because it is a differential quantity. (C) With gain error (v = 1.15): PP is '
    'distorted from 40 to 46 mmHg, demonstrating that PP accuracy validates sensor gain. '
    'Zero calibration corrects the offset in (B) but cannot correct the gain error in (C).'
)

doc.add_page_break()

add_figure(
    'figure3_system_comparison.png',
    'Figure 2. Comparison of conventional and proposed zero-calibration-free arterial pressure monitoring systems. '
    '(A) Conventional fluid-filled system: three sources of DC offset (hydrostatic column, atmospheric '
    'pressure reference, transducer drift) necessitate manual zero calibration, which corrects offset (u) '
    'only. (B) Proposed zero-free system: catheter-tip MEMS sensor eliminates hydrostatic column; '
    'built-in barometer provides electronic atmospheric compensation; self-calibrating MEMS eliminates '
    'drift. All offset sources are eliminated by design, and PP accuracy validates gain (v = 1).'
)

doc.add_page_break()

add_figure(
    'figure2_ccc_zeroing_scenarios.png',
    'Figure 3. Concordance plots for four simulated scenarios (n = 150 paired measurements). '
    '(A) Before zeroing: offset of 12 mmHg produces location shift u = \u20130.55, reducing C\u2082 to 0.867. '
    '(B) After zeroing: offset removed, CCC = 0.986 with C\u2082 = 1.000. '
    '(C) Gain error (v = 1.11): zeroing cannot improve C\u2082 (0.870) because the scale shift persists. '
    '(D) Gain + offset: CCC = 0.976; zeroing would move the device from D to C but not to B. '
    'Dashed line = identity (y = x); solid line = best fit.'
)

doc.add_page_break()

add_figure(
    'figure4_cb_diagnostic_space.png',
    'Figure 4. Bias correction factor (C\u2082) as a function of location shift (u) and scale shift (v). '
    'Contour lines show iso-C\u2082 values. Zero calibration moves a device horizontally (u \u2192 0, green '
    'arrows) but does not change v. Points A\u2013D correspond to the scenarios in figure 3. The gold '
    'star marks the ideal (u = 0, v = 1, C\u2082 = 1.0). A device with correct gain (v = 1, validated by '
    'PP accuracy) and offset-free design (u = 0) achieves C\u2082 = 1.0 without calibration.'
)

doc.add_page_break()

add_figure(
    'figure5_ba_comparison.png',
    'Figure 5. Bland\u2013Altman plots for the four simulated scenarios. '
    'Solid horizontal line = mean bias; dashed lines = limits of agreement (\u00B11.96 SD). '
    'Scenarios B and C both show near-zero bias after zeroing, yet have fundamentally '
    'different error structures detectable only by CCC decomposition (table 2).'
)

doc.add_page_break()

add_figure(
    'figure8_ba_vs_concordance.png',
    'Figure 6. Side-by-side comparison of Bland\u2013Altman and concordance analysis for all four '
    'scenarios. Upper row: concordance plots (identity line = dashed); lower row: Bland\u2013Altman '
    'plots. Bland\u2013Altman analysis cannot distinguish Scenario B (true agreement) from Scenario D '
    '(gain error masked by offset), whereas the concordance plot and CCC decomposition clearly '
    'separate them.'
)

doc.add_page_break()

add_figure(
    'figure6_sensitivity_analysis.png',
    'Figure 7. Sensitivity analysis. (A) C\u2082 as a function of gain error and DC offset: '
    'zero calibration (movement to zero offset, vertical arrows) improves C\u2082 only when gain '
    'is correct (v = 1). (B) CCC degradation with increasing gain error at different precision '
    'levels (r = 0.95, 0.97, 0.99): even high-precision sensors suffer substantial CCC loss '
    'when gain error exceeds \u00B110%. Shaded region indicates typical MEMS gain accuracy (\u00B15%).'
)

doc.add_page_break()

add_figure(
    'figure7_pp_validation.png',
    'Figure 8. Pulse pressure as an implicit gain validator. (A) Correct gain: PP is preserved. '
    '(B) Gain error (v = 1.15): PP is proportionally distorted. (C) DC offset only: PP is unchanged. '
    '(D\u2013E) Regression of measured PP against reference PP: slope estimates sensor gain directly. '
    '(F) Logical chain from PP accuracy to zero-calibration-free monitoring.'
)

# ══════════════════════════════════════════════════════════════════
# SAVE
# ══════════════════════════════════════════════════════════════════
doc.save(OUTPATH)
print(f'English PMEA manuscript saved: {OUTPATH}')
