#!/usr/bin/env python3
"""
Create Physiological Measurement Research Paper .docx — Japanese version

「ゼロ校正で治せないもの：一致性解析がBland–Altmanでは見えない
  ゲインエラーを暴く」

PMEA要件:
  - 構造化抄録: Objective / Approach / Main results / Significance (<=250語)
  - 参考文献: Harvard（著者名アルファベット順）
  - Research paper: <=8000語
"""

import os
import re
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
FIGDIR = os.path.join(SCRIPT_DIR, '..', 'figures')
OUTDIR = os.path.join(SCRIPT_DIR, '..', 'manuscripts')
OUTPATH = os.path.join(OUTDIR, 'PMEA_ZeroFree_Manuscript_JA.docx')

os.makedirs(OUTDIR, exist_ok=True)

doc = Document()

# ── Page setup ──
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(11)
style.paragraph_format.line_spacing = 2.0
style.paragraph_format.space_after = Pt(0)

# ── Helper functions ──
def add_heading_styled(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = 'Times New Roman'
        run.font.color.rgb = RGBColor(0, 0, 0)
    return h

def add_para(text, bold=False, italic=False, alignment=None, space_after=Pt(6)):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(11)
    run.bold = bold
    run.italic = italic
    if alignment:
        p.alignment = alignment
    p.paragraph_format.space_after = space_after
    return p

def add_figure(filename, caption, width=Inches(6)):
    path = os.path.join(FIGDIR, filename)
    if not os.path.exists(path):
        add_para(f'[図ファイルが見つかりません: {filename}]', italic=True)
        cap = add_para(caption, italic=True)
        cap.paragraph_format.space_after = Pt(12)
        return cap
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(path, width=width)
    cap = doc.add_paragraph()
    cap_run = cap.add_run(caption)
    cap_run.font.name = 'Times New Roman'
    cap_run.font.size = Pt(10)
    cap_run.italic = True
    cap.paragraph_format.space_after = Pt(12)
    return cap

# ══════════════════════════════════════════════════════════════════
# 表紙
# ══════════════════════════════════════════════════════════════════
add_para('RESEARCH PAPER（日本語版）', bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run(
    'ゼロ校正で治せないもの：\n'
    '一致性解析がBland\u2013Altmanでは見えないゲインエラーを暴く'
)
run.font.name = 'Times New Roman'
run.font.size = Pt(16)
run.bold = True
title.paragraph_format.space_after = Pt(12)

title_en = doc.add_paragraph()
title_en.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_en = title_en.add_run(
    'What Zeroing Cannot Fix: Concordance Analysis Unmasks\n'
    'Gain Errors Invisible to Bland\u2013Altman'
)
run_en.font.name = 'Times New Roman'
run_en.font.size = Pt(12)
run_en.italic = True
title_en.paragraph_format.space_after = Pt(18)

add_para('ランニングタイトル：ゼロ校正で治せないもの', italic=True,
         alignment=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_paragraph()

add_para('[著者名（未定）]', alignment=WD_ALIGN_PARAGRAPH.CENTER)
add_para('[所属（未定）]', alignment=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_paragraph()

add_para('責任著者:', bold=True)
add_para('[氏名、住所、メールアドレス、ORCID]')
doc.add_paragraph()

add_para('本文語数: 約4,500語（参考文献、表、図の説明文を除く）')
add_para('表・図: 表2、図8点')
add_para('参考文献: 30件')
doc.add_paragraph()

add_para('キーワード: ', bold=True, space_after=Pt(0))
kw = doc.paragraphs[-1]
kw_run = kw.add_run(
    '観血的動脈圧モニタリング；一致性相関係数；ゼロ校正；脈圧；'
    'MEMSセンサ；バイアス補正係数；方法比較；機器バリデーション'
)
kw_run.font.name = 'Times New Roman'
kw_run.font.size = Pt(11)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# 構造化抄録 — PMEA形式
# ══════════════════════════════════════════════════════════════════
add_heading_styled('抄録（Abstract）', level=1)

add_para('Objective（目的）', bold=True, space_after=Pt(2))
add_para(
    '観血的動脈圧トランスデューサのゼロ校正がオフセット成分のみを補正し、ゲインエラーは未検出のまま'
    '残ることをシミュレーションと解析的分解により示すとともに、これらのエラー種別を識別する補完的'
    'バリデーション指標としてLinの一致性相関係数（CCC）分解を提案する。'
)

add_para('Approach（方法）', bold=True, space_after=Pt(2))
add_para(
    'CCCを精度（ピアソンのr）と正確度（バイアス補正係数C\u2082）に分解し、C\u2082をさらに'
    'ロケーションシフト（u、オフセット）とスケールシフト（v、ゲインエラー）に分解する。'
    '臨床的に関連のある4つの動脈圧測定シナリオ（各n = 150対の測定）をシミュレーションした：'
    'オフセットのみ、ゼロ校正後のオフセット除去、ゲインエラー、およびゲイン＋オフセットの複合。'
    '各シナリオをBland\u2013Altman統計量とCCC分解の両方で解析した。また、手動ゼロ校正を不要にする'
    '3つの工学的解決策\u2014カテーテル先端マイクロ電気機械システム（MEMS）センサ、気圧補償、自己校正MEMS\u2014を特定した。'
)

add_para('Main results（主要結果）', bold=True, space_after=Pt(2))
add_para(
    'ゼロ校正後、ゲインエラーの有無にかかわらずBland\u2013Altman解析はほぼゼロのバイアスを示す'
    '（ゲインエラーなし vs. ありでバイアス：\u22120.1 vs. 1.1 mmHg）。一方、CCC分解はこれらのケースを'
    '明確に識別する：C\u2082 = 1.000（ゲインエラーなし）対 C\u2082 = 0.870（10%ゲインエラー、v = 1.11）。'
    '脈圧の正確性が暗黙的なゲイン検証として機能することを示した。脈圧は直流（DC）オフセットから独立だが、'
    'センサゲインに比例するためである。'
)

add_para('Significance（意義）', bold=True, space_after=Pt(2))
add_para(
    'r、C\u2082、u、vへのCCC分解は、Bland\u2013Altman解析単独では捉えられない動脈圧モニターの性能に関する'
    '診断情報を提供する。本枠組みは次世代センサへの定量的設計目標（校正なしでC\u2082 \u2265 0.99）を提示し、'
    '非侵襲的心拍出量モニターにも自然に拡張される。すべての動脈圧バリデーション研究において、'
    '分解を伴うCCCの報告を推奨する。'
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# 緒言
# ══════════════════════════════════════════════════════════════════
add_heading_styled('1. 緒言（Introduction）', level=1)

add_para(
    '観血的動脈圧モニタリングは、麻酔科および集中治療において最も頻繁に実施される測定のひとつである。'
    '動脈カテーテル挿入後、臨床医は血行動態モニタリングを開始する前に、腋窩中線レベル（phlebostatic '
    'axis）で圧トランスデューサを大気圧にゼロ校正しなければならない（Saugel et al 2020, Saugel and '
    'Sessler 2021）。この校正操作は、トランスデューサとカテーテル挿入部位の高さ関係が変化するたびに'
    '\u2014例えば手術台が傾斜した場合やベッドの角度が調整された場合\u2014繰り返す必要がある'
    '（Gupta et al 2025）。圧センシング技術が数十年にわたり進歩してきたにもかかわらず、'
    '手動ゼロ校正の必要性は外部液充填トランスデューサシステムの導入以来、本質的に変わっていない。'
)

add_para(
    'この要件が持続しているのは根本的な物理的必然性ではなく、従来のモニタリングシステムの設計上の帰結'
    'である。標準的な液充填動脈ラインでは、圧トランスデューサは外部に位置し、生理食塩水のカラムによって'
    '血管内カテーテル先端と接続されている。この構造は3つの系統的な直流（DC）オフセット源を導入する：'
    '(1) トランスデューサと測定部位の静水圧差（\u0394P = \u03C1gh、高低差1cmあたり約0.74 mmHg）；'
    '(2) 大気圧（約760 mmHg）に対するゲージ圧測定の必要性；'
    '(3) ストレインゲージの機械的クリープや熱効果によるトランスデューサドリフト'
    '（Mark 1998, McGhee and Bridges 2002）。ゼロ校正はトランスデューサを大気に開放し出力をゼロに'
    'リセットすることで、これら3つのオフセットを同時に除去する。'
)

add_para(
    'しかし、ゼロ校正はもっぱらオフセット補正である。トランスデューサのゲイン（感度）\u2014すなわち'
    '真の圧力変化と電気的出力の比例関係\u2014を検証も補正もしない。ゲインが不正確であれば、測定波形は'
    '拡大または縮小され、このエラーはゼロ校正後も持続する。このオフセットエラーとゲインエラーの区別は'
    '単なる学術的議論ではなく、モニタリング機器のバリデーション方法と、以下で論じるように、'
    '機器設計のあり方に直接的な含意を持つ。'
)

add_para(
    '血行動態モニターの標準的バリデーション枠組みは、Bland\u2013Altman解析'
    '（Bland and Altman 1986, 1999）、Critchley\u2013Critchleyのpercentage error基準'
    '（Critchley and Critchley 1999）、およびトレンド能力評価のためのポーラープロット'
    '（Critchley et al 2010, 2011）を中心とし、参照手技の精度に特に注意が払われている'
    '（Cecconi et al 2009）。これらのツールはバイアス、一致の限界、'
    '方向的一致を効果的に評価するが、いずれも完全一致線に対する正確度と精度を同時に捉える'
    '統合的な単一値指標を提供しない。'
)

add_para(
    '本論文では、機器が脈圧（PP）を正確に測定するならば、そのセンサゲインは暗黙的に検証され、'
    '適切なセンサ設計によりゼロ校正が不要になると提案する。この議論をLinの一致性相関係数（CCC）の'
    '分解（Lin 1989, 2000）を用いて定式化し、各DCオフセット源を排除する工学的解決策を特定し、'
    'シミュレーションによりCCC分解がBland\u2013Altman解析では隠蔽されるゲインエラーを検出することを'
    '示し、動脈圧モニタリングおよびその先への含意を議論する。'
)

# ══════════════════════════════════════════════════════════════════
# 理論と方法
# ══════════════════════════════════════════════════════════════════
add_heading_styled('2. 理論と方法（Theory and methods）', level=1)

add_heading_styled('2.1. 動脈圧の交流・直流成分分解', level=2)

add_para(
    '動脈圧波形P(t)は、緩やかに変化する成分    （直流（DC）レベル：平均動脈圧と外部オフセットにより決定）と拍動成分'
        '（交流（AC）信号：心駆出によって駆動）に分解できる。脈圧（PP = 収縮期血圧[SBP] \u2013 拡張期血圧[DBP]）は純粋なAC量であり、'
    '拍動の最大値と最小値の差を表し、定義上、いかなる加算的DCオフセットからも独立である'
    '（図1、パネルAおよびB）。'
)

add_para(
    'この独立性には重要な含意がある。センサがPPを正確に測定しているならば、そのゲイン（変換感度、'
    'mV/mmHg等の単位）は正しくなければならない。なぜならPPの正確性は、真の圧力差（SBP \u2013 DBP）に'
    '対応する電気的出力差が正確であることを要求するからである。逆に、ゲインが不正確であれば（例えば'
    'センサが1.00 mV/mmHgではなく1.15 mV/mmHgを読む場合）、測定されたPPは比例的に歪む'
    '（図1、パネルC）。したがって、PPの正確性はセンサゲインの暗黙的な妥当性検証として機能する。'
)

add_heading_styled('2.2. DCオフセットの3つの源と工学的排除', level=2)

add_para(
    'ゲインが正しければ（PPの正確性により検証）、残る測定誤差の唯一の源はDCオフセットである。'
    '従来のシステムでは3つの異なる源がこのオフセットに寄与する（表1）。重要なことに、'
    '各々は確立された工学的解決策により排除可能である：'
)

add_para(
    '(1) 静水圧カラムオフセット：カテーテル先端マイクロ電気機械システム（MEMS）圧センサ（例：Millar Mikro-Cath）はセンシング素子を'
    'カテーテル先端に直接配置し、液体カラムを完全に排除する（Hasenkamp et al 2012, Song et al 2020）。'
    'センサは測定対象点で直接圧力を測定し、トランスデューサ\u2013患者間の高低差から独立した測定となる'
    '（Millar 2026）。'
)

add_para(
    '(2) 大気圧基準：絶対圧センサは大気に対してではなく内部真空参照に対して全圧を測定する。モニター内に'
    '気圧センサを組み込み、大気圧を電子的に減算することで、手動の大気圧ゼロ校正なしにゲージ圧を算出できる。'
    'このアプローチはMillar社のTiSenseプラットフォームなどの慢性埋込型センサにすでに実装されている'
    '（Millar 2026b）。'
)

add_para(
    '(3) トランスデューサドリフト：自己校正MEMSセンサは、周期的な自動校正のための既知の圧力点を提供する'
    '内部参照圧力キャビティを搭載している（Kang et al 2022）。密封マイクロキャビティ内の液体\u2013気体'
    '相転移を利用して参照圧力を巡回することで、センサは外部介入なしにゼロドリフトを継続的に補正する。'
)

add_para(
    '3つの解決策すべてを同時に実装すれば\u2014先端センサ、気圧補償付き絶対圧、および自己校正参照\u2014'
    'すべてのDCオフセット源が設計により排除される。手動ゼロ校正は冗長となる（図2）。'
)

# ══════════════════════════════════════════════════════════════════
# CCC枠組み
# ══════════════════════════════════════════════════════════════════
add_heading_styled('2.3. LinのCCCとその分解', level=2)

add_para(
    'Linの一致性相関係数（\u03C1c）は、対をなす測定値の45度完全一致線上の合致度を定量化する'
    '（Lin 1989）。以下のように分解される：'
)

eq = doc.add_paragraph()
eq.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = eq.add_run('\u03C1c = r \u00D7 C\u2082')
run.font.name = 'Times New Roman'
run.font.size = Pt(12)
run.italic = True
eq.paragraph_format.space_after = Pt(6)

add_para(
    'ここでr（ピアソンの相関係数）は精度（precision：回帰直線周りのデータの緊密さ）を測定し、'
    'C\u2082（バイアス補正係数）は正確度（accuracy：回帰直線の45度恒等線からの乖離）を測定する。'
    'C\u2082はさらに以下のように分解される（Lin 2000）：'
)

eq2 = doc.add_paragraph()
eq2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = eq2.add_run('C\u2082 = 2 / (v + 1/v + u\u00B2)')
run2.font.name = 'Times New Roman'
run2.font.size = Pt(12)
run2.italic = True
eq2.paragraph_format.space_after = Pt(6)

add_para(
    'ここでv = \u03C31/\u03C32はスケールシフト（標準偏差の比、ゲインエラーを反映）、'
    'u = (\u03BC1 \u2013 \u03BC2)/\u221A(\u03C31\u03C32)はロケーションシフト（正規化された平均差、'
    'オフセットを反映）である。'
)

add_heading_styled('2.4. CCC的観点からゼロ校正が補正するもの・しないもの', level=2)

add_para(
    'ゼロ校正は、機器の読み値と真の圧力の間の系統的オフセットを除去することでuをゼロに近づける操作である。'
    'CCC的には、ゼロ校正の成功はu \u2248 0を達成し、ロケーション成分に関してC\u2082を最大化する。'
    'しかし、ゼロ校正はvに影響しない：センサゲインが不正確であれば（v \u2260 1）、完全なゼロ校正後も'
    'C\u2082は1.0未満のままである（図3、図4）。'
)

add_heading_styled('2.5. シミュレーション設計', level=2)

add_para(
    '臨床的に関連のある4つの動脈圧測定シナリオをシミュレーションした。各シナリオはn = 150対の測定で構成し、'
    '真の収縮期血圧は80〜180 mmHgの範囲とした。'
    'シナリオA：オフセットのみ（ゼロ校正前に12 mmHgの系統的オフセット）。'
    'シナリオB：ゼロ校正後（オフセット除去、ゲインエラーなし）。'
    'シナリオC：ゲインエラーのみ（v = 1.11、すなわち圧力変化の11%過大評価）。'
    'シナリオD：ゲインエラー＋オフセットの複合。'
    '生理的・測定的変動をシミュレートするため、すべてのシナリオにガウスノイズ（標準偏差[SD] = 3.5 mmHg）を加えた。'
    '各シナリオをBland\u2013Altman統計量（バイアス、一致の限界、percentage error）と'
    'CCC分解（\u03C1c、r、C\u2082、u、v）の両方で解析した。'
)

add_heading_styled('2.6. 完全な論証', level=2)

add_para(
    '工学的観点と統計的観点を統合すると、以下のように形式的に述べることができる：'
)

add_para('(1) PPが正確 \u2192 ゲインが正しい \u2192 v = 1（スケールシフトなし）')
add_para('(2) 先端センサ＋気圧補償＋自己校正MEMS \u2192 全DCオフセット排除 \u2192 u = 0（ロケーションシフトなし）')
add_para('(3) v = 1 かつ u = 0 \u2192 C\u2082 = 1.0')
add_para('(4) CCC = r \u00D7 1.0 = r')

add_para(
    'すなわち、適切に設計されたセンサの性能は精度（r）のみによって制限される。rはランダムな測定ノイズを'
    '反映する。すべての系統的誤差\u2014オフセットとゲインの両方\u2014は校正ではなく設計により解決される。'
    'センサのC\u2082は校正の結果ではなく、設計仕様として1.0に近づく。'
)

# ══════════════════════════════════════════════════════════════════
# 結果
# ══════════════════════════════════════════════════════════════════
add_heading_styled('3. 結果（Results）', level=1)

add_heading_styled('3.1. シミュレーション結果', level=2)

add_para(
    '図3に4つのシミュレーションシナリオのコンコーダンスプロットを示す。'
    'シナリオA（オフセットのみ、ゼロ校正前）はロケーションシフト（u = \u20130.55）によりCCCが低下'
    '（0.855）している。シナリオB（ゼロ校正後）はオフセットが除去され、ゲインエラーがないため'
    'ほぼ完全な一致（CCC = 0.986、C\u2082 = 1.000）を示す。シナリオC（ゲインエラー、v = 1.11）は、'
    'スケールシフトが持続するためゼロ校正ではC\u2082（0.870）を改善できないことを示す。'
    'シナリオD（ゲイン＋オフセット）はCCC = 0.976であり、ゼロ校正はDからCへの改善は'
    'もたらすが、Bへの到達はできない。'
)

add_heading_styled('3.2. Bland\u2013Altmanとの比較', level=2)

add_para(
    '重要な比較はシナリオBとCの間にある：Bland\u2013Altman解析（図5、図6）はいずれの場合も'
    'バイアスがほぼゼロであることを示すが、CCC分解は根本的に異なるエラー構造を明らかにする（表2）。'
    'シナリオBはC\u2082 = 1.000（系統的エラーなし）であるのに対し、シナリオCはC\u2082 = 0.870'
    '（隠れたゲインエラー、v = 1.11）である。これはBland\u2013Altman解析のみでは、真の一致と、'
    'ゲインエラーを隠蔽するオフセット相殺による一致を区別できないことを実証している。'
)

add_heading_styled('3.3. 感度分析', level=2)

add_para(
    '図7はこれらの関係を定量化する。パネルAはゲインエラーとDCオフセットの関数としてのC\u2082を示し、'
    'ゼロ校正（ゼロオフセットへの移動）はゲインが正しい場合にのみC\u2082を改善することを確認する。'
    'パネルBは異なるセンサ精度（r）レベルでのゲインエラーに伴うCCCの劣化を示し、高精度センサ'
    '（r = 0.99）であってもゲインエラーが\u00B110%を超えるとCCCが大幅に低下することを実証する。'
    '特筆すべきは、典型的なMEMSセンサは\u00B15%以内のゲイン精度を達成しており、'
    'ゼロオフセット条件でC\u2082 > 0.99に相当することである。'
)

add_heading_styled('3.4. ゲイン検証としての脈圧', level=2)

add_para(
    '図8はPP\u2013ゲイン関係の定量的実証を提供する。正しいゲイン（パネルA）、ゲインエラー（パネルB）、'
    'DCオフセットのみ（パネルC）のシミュレーション動脈圧波形は、PPがゲインエラーにより歪むが'
    'DCオフセットには影響されないことを示す。パネルDおよびEは、基準に対するPPの回帰直線の傾きが'
    'センサゲインを直接推定することを確認し、実用的なバリデーションツールを提供する。'
    'PPの正確性からゼロ校正不要モニタリングへの論理的連鎖をパネルFに要約する。'
)

# ══════════════════════════════════════════════════════════════════
# 考察
# ══════════════════════════════════════════════════════════════════
add_heading_styled('4. 考察（Discussion）', level=1)

add_heading_styled('4.1. センサ設計への含意', level=2)

add_para(
    '本論文で提示した枠組みは、ゼロ校正を臨床的必要性から設計上の回避策へと再定義する。従来の液充填'
    'システムは、その構造が本質的にDCオフセットを導入するためにゼロ校正を必要とする。校正手順を洗練し'
    '続けるのではなく、機器メーカーは校正の必要性そのものを排除する設計を追求すべきである。'
    '構成技術\u2014カテーテル先端MEMS、絶対圧センシング、気圧補償、および自己校正参照\u2014は、'
        '商用または商用化間近の製品に個別にすでに存在しており（Hasenkamp et al 2012, Song et al 2020, '
        'Millar 2026, 2026b, Kang et al 2022）、高忠実度圧力ワイヤーがカテーテルベースの正確な圧測定の'
        '実現可能性をすでに実証している（Scalia et al 2023）。これらを単一の臨床用動脈圧モニタリングシステム'
        'に統合することは、科学的課題ではなく工学的課題である。'
)

add_para(
    'CCC分解は定量的な設計目標を提供する：機器開発者は、校正ステップなしでC\u2082 \u2265 0.99を達成すること'
    'を目指すべきであり\u2014McBrideの提案する一致強度基準（McBride 2005）における「卓越した」一致に相当する'
    '\u2014これは未校正の機器出力を基準標準と比較することで検証可能である。校正前に'
    'C\u2082 < 0.99であれば、機器には校正が部分的にしか隠蔽できない残存系統的エラーがある。'
    'uとvへの分解により、残存エラーがオフセット（ゼロ校正で対処可能）かゲインエラー'
    '（ハードウェアまたはアルゴリズムの修正が必要）かをさらに特定できる。'
)

add_heading_styled('4.2. バリデーション方法論への含意', level=2)

add_para(
    '現在の規制経路（例：米国食品医薬品局[FDA] 510(k)）は、動脈圧モニターのバリデーションに特定の統計手法を規定して'
        'いない（FDA 2026）。公表されたバリデーション研究は、侵襲的・非侵襲的双方の動脈圧モニターを含め、'
        '圧倒的にBland\u2013Altman解析とpercentage errorに'
        '依拠している（Kim et al 2014, Joosten et al 2017, Bland and Altman 1986）。しかし、ゼロ校正後の機器の'
    'Bland\u2013Altmanプロットはバイアス \u2248 0を示し、一致の限界内に比例バイアス（ゲインエラー）を'
    '隠蔽する可能性がある。CCC報告、特にC\u2082成分は、オフセット補正と真の測定精度を区別する追加の'
    '精査層を提供する。'
)

add_para(
    '我々は、動脈圧モニターのバリデーション研究が以下を報告すべきことを提案する：(1) rとC\u2082への分解を'
    '伴うCCC；(2) ゼロ校正前後のC\u2082（機器の校正依存度を定量化）；(3) C\u2082低下へのuとvの個別の寄与。'
    'この情報により、規制当局と臨床医は、機器が良好な設計によって正確性を達成しているのか、'
    '校正依存的なオフセット除去によって達成しているのかを評価できるようになる。'
)

add_heading_styled('4.3. 非侵襲的心拍出量モニターへの拡張', level=2)

add_para(
    '校正がオフセット（u）のみを補正し、ゲインエラー（v）には影響しないという原理は、非侵襲的心拍出量'
    '（CO）モニターにも自然に拡張される。COモニタリングにおいては、肺動脈カテーテルが臨床参照標準として'
    '依然として重要な位置を占める（Chatterjee 2009）。ClearSightのPhysiocalアルゴリズムはVolume clampのセットポイント'
    'を周期的に再最適化する\u2014本質的にオフセット補正である（Ameloot et al 2015）。'
    '非侵襲的心拍出量モニタリング（NICOM）/Starlingのバイオリアクタンス位相基準は位相\u2013一回拍出量変換のベースラインオフセットを提供する'
    '（Squara et al 2007）。FloTracの動脈圧波形較正は平均圧オフセットを調整する（Manecke 2005, Romagnoli et al 2013）。'
    'いずれの場合も、自動校正ルーチンはu（ロケーションシフト）を処理するが、v（スケールシフト）'
    '\u2014すなわち圧力\u2192COまたはインピーダンス\u2192一回拍出量（SV）の変換ゲイン\u2014は補正しない。'
)

add_para(
    'この観察は、Odorら（2017）が提案したCOモニタリングバリデーション枠組みおよび血行動態モニタリングにおける'
    'CCC採用の提唱と組み合わせることで、統一原理を示唆する：いかなる血行動態モニタリング機器においても、'
    'バイアス補正係数C\u2082をuとvの成分に分解し、許容可能なC\u2082を達成するための機器の校正依存度を'
    '明示的に定量化すべきである。C\u2082を1.0近くに維持するために頻繁な再校正を必要とする機器は、校正が部分的'
    'にしか補償できない根本的な設計上の限界を抱えている。'
)

add_heading_styled('4.4. 前提条件と限界', level=2)

add_para(
    '本論文の議論は、センサ応答が生理的圧力範囲にわたって線形であるという前提に基づいている。センサが '
    '有意な非線形性を示す場合、ある圧力レベルでのPP精度は他のレベルでのゲイン正確性を保証しない。'
    'しかし、現代のMEMSピエゾ抵抗素子は一般に0\u2013300 mmHgにわたり全スケール出力の0.1%未満の'
    '非線形性を達成しており（Barlian et al 2009）、臨床的動脈圧モニタリングの要件を十分に満たしている。'
)

add_para(
    '本枠組みは系統的誤差のみを扱う。ランダム測定ノイズ（rを決定する）は別の問題であり、校正や '
    'ここで議論した設計特性の影響を受けない。高いランダムノイズを持つゼロ校正不要センサは、'
    'C\u2082 = 1.0を達成しても依然として低いCCCを示す。'
)

add_para(
    '最後に、日常的な動脈圧モニタリングにカテーテル先端MEMSセンサを臨床採用することには、コスト、'
    'ディスポーザビリティ、および既存モニタリングインフラとの互換性を含む実用的障壁があることを認める。'
    '我々の目的は即座の臨床変更を強制することではなく、ゼロ校正が設計により除去可能な限界である'
    'という理論的原理を確立し、この目標への進捗を評価する定量的枠組み（CCC分解）を提供することである。'
)

add_heading_styled('4.5. 臨床的視点', level=2)

add_para(
    '臨床医にとっての即座の実践的メッセージは微妙なものである。現行の液充填システムではゼロ校正は '
    '依然として不可欠であり、Saugelら（2020）やGuptaら（2025）が強調するように厳密に実施し続けるべきで'
    'ある。しかし、臨床医は完全なゼロ校正であってもオフセットのみを補正し、ゲインは補正しないことを'
    '認識すべきである。10%のゲインエラーを持つゼロ校正済みトランスデューサは、SBPを120ではなく132、'
    'DBPを80ではなく88と表示する\u2014ゼロ校正手順では検出されず、ベースラインのシフトではなく'
    '拡大された脈圧（40 vs. 52 mmHg）として現れる、臨床的に有意な過大評価である。'
)

add_para(
    'PPの正確性がゲインの正しさを暗示するという認識は、簡便な臨床的チェックも示唆する：侵襲的に測定された'
    'PPが生理学的に妥当であり、非侵襲的カフ測定と一致していれば、センサゲインはおそらく正しい。'
    '非生理学的なPP（例えば患者の状態に対して非常に狭いまたは非常に広い）は、ゼロ校正では解決できない'
    'ゲインの問題を示唆する可能性がある。'
)

# ══════════════════════════════════════════════════════════════════
# 結論
# ══════════════════════════════════════════════════════════════════
add_heading_styled('5. 結論（Conclusion）', level=1)

add_para(
    '観血的動脈圧トランスデューサのゼロ校正がオフセット（ロケーションシフト、u）のみを補正し、'
    'ゲインエラー（スケールシフト、v）は未検出のままであることを示すシミュレーション枠組みを提示した。'
    '脈圧の正確性がセンサゲインを検証し、スケールシフトを排除する（v = 1）。カテーテル先端MEMSセンサ、'
    '気圧補償付き絶対圧測定、および自己校正MEMS参照が、すべてのDCオフセット源を排除する（u = 0）。'
    'LinのCCC分解を通じて定式化すると、C\u2082 = 1.0は校正ではなく設計により達成可能であることを意味する。'
    '本枠組みは、次世代動脈圧モニターの定量的設計目標、規制上のバリデーションへの補完的指標、'
    'および生理的測定における校正が何を行い\u2014何を行わないか\u2014のより深い理解を提供する。'
)

# ══════════════════════════════════════════════════════════════════
# 宣言
# ══════════════════════════════════════════════════════════════════
add_heading_styled('利益相反の開示', level=1)
add_para('[著者が記入]')

add_heading_styled('資金', level=1)
add_para('[著者が記入]')

add_heading_styled('著者の貢献', level=1)
add_para('[著者が記入]')

add_heading_styled('データの利用可能性', level=1)
add_para(
    '本論文の図表および結果の生成に使用したすべてのシミュレーションコードとデータは、'
    '合理的な要請に応じて責任著者より入手可能である。'
)

add_heading_styled('謝辞', level=1)
add_para('[著者が記入]')

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# 参考文献 — Harvard（著者名アルファベット順）
# ══════════════════════════════════════════════════════════════════
add_heading_styled('参考文献（References）', level=1)

refs_harvard = [
    'Ameloot K, Palmers P J and Malbrain M L N G 2015 The accuracy of noninvasive cardiac output and pressure measurements with finger cuff: a concise review Curr. Opin. Crit. Care 21 232\u20139',
    'Barlian A A, Park W-T, Mallon J R, Rastegar A J and Pruitt B L 2009 Review: semiconductor piezoresistance for microsystems Proc. IEEE 97 513\u201352',
    'Bland J M and Altman D G 1986 Statistical methods for assessing agreement between two methods of clinical measurement Lancet 327 307\u201310',
    'Bland J M and Altman D G 1999 Measuring agreement in method comparison studies Stat. Methods Med. Res. 8 135\u201360',
    'Cecconi M, Rhodes A, Poloniecki J, Della Rocca G and Grounds R M 2009 Bench-to-bedside review: the importance of the precision of the reference technique in method comparison studies Crit. Care 13 201',
    'Chatterjee K 2009 The Swan\u2013Ganz catheters: past, present, and future Circulation 119 147\u201352',
    'Critchley L A H and Critchley J A J H 1999 A meta-analysis of studies using bias and precision statistics to compare cardiac output measurement techniques J. Clin. Monit. Comput. 15 85\u201391',
    'Critchley L A, Lee A and Ho A M-H 2010 A critical review of the ability of continuous cardiac output monitors to measure trends in cardiac output Anesth. Analg. 111 1180\u201392',
    'Critchley L A, Yang X X and Lee A 2011 Assessment of trending ability of cardiac output monitors by polar plot methodology J. Cardiothorac. Vasc. Anesth. 25 536\u201346',
    'FDA (U.S. Food and Drug Administration) 2026 Premarket Notification 510(k) https://www.fda.gov/medical-devices/premarket-submissions-selecting-and-preparing-correct-submission/premarket-notification-510k (accessed 25 March 2026)',
    'Gupta D, Jain A and Ismaeil M 2025 Zero arterial catheters with every change in the height difference of pressure transducer and catheter insertion site Ann. Card. Anaesth. 28 205',
    'Hasenkamp W, Theumer T, Nolte J, Aschenbrenner L, Dietzel A and B\u00FCttgenbach S 2012 Polyimide/SU-8 catheter-tip MEMS gauge pressure sensor Biomed. Microdevices 14 819\u201328',
    'Joosten A, Desebbe O, Suehiro K, Murphy L S-L, Essber H, Alexander B, Fischer M O, Barvais L, Van Obbergh L, Maucort-Boulch D and Cannesson M 2017 Accuracy and precision of non-invasive cardiac output monitoring devices in perioperative medicine: a systematic review and meta-analysis Br. J. Anaesth. 118 298\u2013310',
    'Kang Y, Ge C, Hu Y, Gao J, Zhong M, Du H and Wang S 2022 Development of a flexible integrated self-calibrating MEMS pressure sensor using a liquid-to-vapor phase change Sensors 22 9737',
    'Kim S-H, Lilot M, Sidhu K S, Rinehart J, Cannesson M, Joosten A and Lee C 2014 Accuracy and precision of continuous noninvasive arterial pressure monitoring compared with invasive arterial pressure Anesthesiology 120 1080\u201397',
    'Lin L I-K 1989 A concordance correlation coefficient to evaluate reproducibility Biometrics 45 255\u201368',
    'Lin L I-K 2000 A note on the concordance correlation coefficient Biometrics 56 324\u20135',
    'Manecke G R 2005 Edwards FloTrac sensor and Vigileo monitor: easy, accurate, reliable cardiac output assessment using the arterial pulse wave Expert Rev. Med. Devices 2 523\u20137',
    'Mark J B 1998 Atlas of Cardiovascular Monitoring (New York: Churchill Livingstone)',
    'McBride G B 2005 A proposal for strength-of-agreement criteria for Lin\u2019s concordance correlation coefficient NIWA Client Report HAM2005-062',
    'McGhee B H and Bridges M E J 2002 Monitoring arterial blood pressure: what you may not know Crit. Care Nurse 22 60\u201379',
    'Millar, Inc. 2026 Mikro-Cath Pressure Catheter \u2013 Product Information https://millar.com/Clinical/MikroCath/ (accessed 25 March 2026)',
    'Millar, Inc. 2026b TiSense \u2013 Chronic Pressure Sensing https://millar.com/Clinical/MikroCath/Continuous-Compartment-Pressure-Measurements/ (accessed 25 March 2026)',
    'Odor P M, Bampoe S and Cecconi M 2017 Cardiac output monitoring: validation studies\u2014how results should be presented Curr. Anesthesiol. Rep. 7 410\u201315',
    'Romagnoli S, Ricci Z, Romano S M, Dimizio F, Bonicolini E, Quattrone D, De Gaudio A R and Grasso S 2013 FloTrac/Vigileo (third generation) and MostCare/PRAM versus echocardiography for cardiac output estimation in vascular surgery J. Cardiothorac. Vasc. Anesth. 27 1114\u201321',
    'Saugel B, Kouz K, Meidert A S, Schulte-Uentrop L and Romagnoli S 2020 How to measure blood pressure using an arterial catheter: a systematic 5-step approach Crit. Care 24 172',
    'Saugel B and Sessler D I 2021 Perioperative blood pressure management Anesthesiology 134 250\u201361',
    'Scalia A, Ghafari C, Navarre W, Delmotte P, Phillips R and Carlier S 2023 High fidelity pressure wires provide accurate validation of non-invasive central blood pressure and pulse wave velocity measurements Biomedicines 11 1235',
    'Song P, Ma Z, Ma J, Yang L, Wei J, Zhao Y, Zhang M, Yang F and Wang X 2020 Recent progress of miniature MEMS pressure sensors Micromachines 11 56',
    'Squara P, Denjean D, Estagnasie P, Brusset A, Dib J C and Dubois C 2007 Noninvasive cardiac output monitoring (NICOM): a clinical validation Intensive Care Med. 33 1191\u20134',
]

for ref in refs_harvard:
    p = doc.add_paragraph()
    run = p.add_run(ref)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(10)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.first_line_indent = Cm(-1.27)
    p.paragraph_format.left_indent = Cm(1.27)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# 表1
# ══════════════════════════════════════════════════════════════════
add_heading_styled('表1（Table 1）', level=1)
add_para(
    '表1. 従来の観血的動脈圧モニタリングにおけるDCオフセットの源と排除のための工学的解決策',
    bold=True, italic=True
)

table = doc.add_table(rows=5, cols=5)
table.style = 'Light Grid Accent 1'

headers = ['オフセット源', '物理的機序', '大きさ', '工学的解決策', 'CCC成分への影響']
for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = h
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True
            run.font.size = Pt(9)
            run.font.name = 'Times New Roman'

data = [
    ['静水圧カラム', '\u0394P = \u03C1gh（トランスデューサとカテーテル先端の高低差）',
     '~0.74 mmHg/cm', 'カテーテル先端MEMSセンサ（液体カラムを排除）',
     'uを低減（ロケーションシフト）'],
    ['大気圧基準', '大気（~760 mmHg）に対するゲージ圧測定',
     'ベースライン全体', '絶対圧センサ＋内蔵気圧計による電子的減算',
     'uを低減（ロケーションシフト）'],
    ['トランスデューサドリフト', 'ストレインゲージの機械的クリープ・熱効果',
     '~1\u20135 mmHg/日', '内部参照圧力キャビティ付き自己校正MEMS',
     'uを低減（ロケーションシフト）'],
    ['ゲインエラー（ゼロ校正で補正不可）', '感度不一致：単位圧力あたりの出力 \u2260 公称値',
     '1\u201315%（典型的）', '工場校正；PPの正確性による暗黙的ゲイン検証',
     'vに影響（スケールシフト）；ゼロ校正は無効'],
]

for row_idx, row_data in enumerate(data, start=1):
    for col_idx, cell_text in enumerate(row_data):
        cell = table.rows[row_idx].cells[col_idx]
        cell.text = cell_text
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(8)
                run.font.name = 'Times New Roman'

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# 表2
# ══════════════════════════════════════════════════════════════════
add_heading_styled('表2（Table 2）', level=1)
add_para(
    '表2. 4つのシミュレーションシナリオにおける統計指標の比較（n = 150対の測定）',
    bold=True, italic=True
)

table2 = doc.add_table(rows=5, cols=10)
table2.style = 'Light Grid Accent 1'

t2_headers = ['シナリオ', 'CCC', 'r', 'C\u2082', 'u', 'v', 'バイアス\n(mmHg)', 'LoA下限\n(mmHg)',
              'LoA上限\n(mmHg)', 'PE (%)']
for i, h in enumerate(t2_headers):
    cell = table2.rows[0].cells[i]
    cell.text = h
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True
            run.font.size = Pt(8)
            run.font.name = 'Times New Roman'

t2_data = [
    ['A: ゼロ校正前', '0.855', '0.986', '0.867', '\u20130.55', '0.99', '11.9', '4.8', '19.0', '7.2'],
    ['B: ゼロ校正後', '0.986', '0.986', '1.000', '0.01', '0.99', '\u20130.1', '\u20137.2', '7.0', '7.2'],
    ['C: ゲインエラー', '0.855', '0.982', '0.870', '0.54', '1.11', '\u201310.9', '\u201319.4', '\u20132.4', '8.7'],
    ['D: ゲイン＋オフセット', '0.976', '0.982', '0.993', '\u20130.05', '1.11', '1.1', '\u20137.4', '9.6', '8.7'],
]

for row_idx, row_data in enumerate(t2_data, start=1):
    for col_idx, cell_text in enumerate(row_data):
        cell = table2.rows[row_idx].cells[col_idx]
        cell.text = cell_text
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(8)
                run.font.name = 'Times New Roman'

add_para(
    'CCC = Linの一致性相関係数；r = ピアソンの相関係数（精度）；C\u2082 = バイアス補正係数（正確度）；'
    'u = ロケーションシフト（正規化されたオフセット）；v = スケールシフト（ゲイン比）；'
    'LoA = 一致の限界（Bland\u2013Altman）；PE = percentage error。'
    'シナリオBとCは類似したバイアス（ほぼ0）とLoAを示すが、C\u2082は大幅に異なり（1.000 vs. 0.870）、'
    'Bland\u2013Altman解析のみではゲインエラーを検出できないことを実証している。',
    italic=True
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# 図
# ══════════════════════════════════════════════════════════════════
add_heading_styled('図（Figures）', level=1)

add_figure(
    'figure1_signal_decomposition.png',
    '図1. 動脈圧波形の分解。(A) 真の動脈圧：DC成分（オフセット依存のベースライン）とAC成分'
    '（脈圧、PP = 40 mmHg）。(B) 15 mmHgのDCオフセット（例：静水圧カラム）あり：PPは差分量であるため'
    '40 mmHgのまま不変。(C) ゲインエラー（v = 1.15）あり：PPが40から46 mmHgに歪み、'
    'PPの正確性がセンサゲインを検証することを実証。ゼロ校正は(B)のオフセットを補正するが、'
    '(C)のゲインエラーは補正できない。'
)

doc.add_page_break()

add_figure(
    'figure3_system_comparison.png',
    '図2. 従来型と提案するゼロ校正不要動脈圧モニタリングシステムの比較。'
    '(A) 従来の液充填システム：3つのDCオフセット源（静水圧カラム、大気圧基準、トランスデューサドリフト）'
    'が手動ゼロ校正を必要とし、校正はオフセット（u）のみを補正する。'
    '(B) 提案するゼロ校正不要システム：カテーテル先端MEMSセンサが静水圧カラムを排除、内蔵気圧計が'
    '大気圧を電子的に補償、自己校正MEMSがドリフトを排除。すべてのオフセット源が設計により排除され、'
    'PPの正確性がゲインを検証する（v = 1）。'
)

doc.add_page_break()

add_figure(
    'figure2_ccc_zeroing_scenarios.png',
    '図3. 4つのシミュレーションシナリオのコンコーダンスプロット（n = 150対の測定）。'
    '(A) ゼロ校正前：12 mmHgのオフセットによりu = \u20130.55のロケーションシフトが生じ、C\u2082が0.867に低下。'
    '(B) ゼロ校正後：オフセット除去、CCC = 0.986、C\u2082 = 1.000。'
    '(C) ゲインエラー（v = 1.11）：スケールシフトが持続するためゼロ校正ではC\u2082（0.870）を改善不可。'
    '(D) ゲイン＋オフセット：CCC = 0.976；ゼロ校正はDからCへの移動のみ、Bには到達不可。'
    '破線 = 恒等線（y = x）；実線 = 回帰直線。'
)

doc.add_page_break()

add_figure(
    'figure4_cb_diagnostic_space.png',
    '図4. ロケーションシフト（u）とスケールシフト（v）の関数としてのバイアス補正係数（C\u2082）。'
    '等高線はiso-C\u2082値を示す。ゼロ校正は機器を水平方向に移動（u \u2192 0、緑色矢印）させるが、'
    'vは変化しない。点A\u2013Dは図3のシナリオに対応する。金色の星は理想（u = 0, v = 1, C\u2082 = 1.0）'
    'を示す。正しいゲイン（v = 1、PPの正確性により検証）とオフセットフリーの設計（u = 0）を持つ機器は、'
    '校正なしでC\u2082 = 1.0を達成する。'
)

doc.add_page_break()

add_figure(
    'figure5_ba_comparison.png',
    '図5. 4つのシミュレーションシナリオのBland\u2013Altmanプロット。'
    '実線 = 平均バイアス；破線 = 95%一致の限界（\u00B11.96 SD）。'
    'シナリオBとCはいずれもゼロ校正後にバイアスほぼゼロを示すが、'
    'CCC分解（表2）によってのみ検出可能な根本的に異なるエラー構造を持つ。'
)

doc.add_page_break()

add_figure(
    'figure8_ba_vs_concordance.png',
    '図6. コンコーダンスプロット（上段）とBland\u2013Altmanプロット（下段）の並列比較。'
    '4つのシナリオ（A\u2013D）について同一データを2つの方法で可視化。'
    'Bland\u2013Altman解析はシナリオB（真の一致）とシナリオD（ゲインエラーがオフセットにより隠蔽）を'
    '区別できないが、コンコーダンスプロットとCCC分解はこれらを明確に分離する。'
)

doc.add_page_break()

add_figure(
    'figure6_sensitivity_analysis.png',
    '図7. 感度分析。(A) ゲインエラー（%）とDCオフセット（mmHg）の関数としてのC\u2082値：'
    'ゼロ校正（ゼロオフセットへの移動、垂直矢印）はゲインが正しい場合にのみC\u2082を改善する。'
    '(B) 異なるセンサ精度レベル（r = 0.95, 0.97, 0.99）でのゲインエラーに伴うCCC劣化：'
    '高精度センサでもゲインエラー > \u00B110%でCCCが大幅に低下。'
    '灰色帯域 = 典型的なMEMSゲイン精度範囲（\u00B15%）。'
)

doc.add_page_break()

add_figure(
    'figure7_pp_validation.png',
    '図8. ゲイン検証としての脈圧（PP）。'
    '(A) 正しいゲイン：PPは保存される。'
    '(B) ゲインエラー（v = 1.15）：PPが比例的に歪む。'
    '(C) DCオフセットのみ：PPは不変。'
    '(D\u2013E) 基準PPに対する測定PPの回帰：傾きがセンサゲインを直接推定。'
    '(F) PPの正確性からゼロ校正不要モニタリングへの論理連鎖の要約。'
)

# ══════════════════════════════════════════════════════════════════
# SAVE
# ══════════════════════════════════════════════════════════════════
doc.save(OUTPATH)
print(f'Japanese PMEA manuscript saved: {OUTPATH}')
