#!/usr/bin/env python3
"""
Create cover letter .docx for Physiological Measurement submission.
"""

import os
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
OUTDIR = os.path.join(SCRIPT_DIR, '..', 'cover_letter')
OUTPATH = os.path.join(OUTDIR, 'PMEA_CoverLetter_EN.docx')

os.makedirs(OUTDIR, exist_ok=True)

doc = Document()

# Page setup
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(11)
style.paragraph_format.line_spacing = 1.15
style.paragraph_format.space_after = Pt(6)


def add_para(text, bold=False, italic=False, space_after=Pt(6), alignment=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(11)
    run.bold = bold
    run.italic = italic
    if alignment:
        p.alignment = alignment
    p.paragraph_format.space_after = space_after
    return p


# Date
add_para('[Date]', space_after=Pt(12))

# Addressee
add_para('Editor-in-Chief')
p = add_para('Physiological Measurement', italic=True)
add_para('IOP Publishing', space_after=Pt(12))

add_para('Dear Editor,', space_after=Pt(12))

# Opening paragraph
p = doc.add_paragraph()
run = p.add_run('We are pleased to submit our manuscript entitled ')
run.font.name = 'Times New Roman'
run.font.size = Pt(11)
run = p.add_run(
    '"What Zeroing Cannot Fix: Concordance Analysis Unmasks Gain Errors '
    'Invisible to Bland\u2013Altman"'
)
run.font.name = 'Times New Roman'
run.font.size = Pt(11)
run.bold = True
run = p.add_run(' for consideration as a Research Paper in ')
run.font.name = 'Times New Roman'
run.font.size = Pt(11)
run = p.add_run('Physiological Measurement')
run.font.name = 'Times New Roman'
run.font.size = Pt(11)
run.italic = True
run = p.add_run('.')
run.font.name = 'Times New Roman'
run.font.size = Pt(11)
p.paragraph_format.space_after = Pt(12)

# Summary section
add_para('Summary', bold=True, space_after=Pt(4))

add_para(
    'This manuscript presents a simulation-based framework demonstrating that zero calibration '
    'of invasive arterial pressure transducers corrects only the offset component (location '
    'shift, u) while leaving gain errors (scale shift, v) undetected. Using Lin\u2019s concordance '
    'correlation coefficient (CCC) decomposition into precision (r) and accuracy (C\u2082), we show '
    'that Bland\u2013Altman analysis alone cannot distinguish true agreement from apparent agreement '
    'caused by offsetting gain errors. We further identify three engineering solutions \u2014 '
    'catheter-tip MEMS sensors, barometric compensation, and self-calibrating MEMS \u2014 that '
    'eliminate all sources of DC offset by design, rendering manual zero calibration redundant.',
    space_after=Pt(12)
)

# Relevance section
add_para('Relevance to Physiological Measurement', bold=True, space_after=Pt(4))

add_para(
    'We believe this work is well suited to Physiological Measurement for the following reasons:'
)

add_para(
    '1. Measurement methodology and validation: The manuscript directly addresses how '
    'physiological measurements are validated, proposing CCC decomposition as a complementary '
    'metric to Bland\u2013Altman analysis \u2014 a topic central to the journal\u2019s scope.'
)

add_para(
    '2. Sensor design and engineering: We bridge clinical measurement practice with MEMS sensor '
    'engineering, identifying specific design targets (C\u2082 \u2265 0.99 without calibration) that '
    'can guide next-generation arterial pressure monitors.'
)

add_para(
    '3. Quantitative framework: The CCC decomposition framework provides a rigorous, '
    'analytically grounded approach to understanding what calibration does \u2014 and does not \u2014 '
    'correct, with direct implications for regulatory validation of haemodynamic monitors.'
)

add_para(
    '4. Broader applicability: The framework extends naturally to non-invasive cardiac output '
    'monitors (ClearSight, NICOM/Starling, FloTrac), offering a unified principle for evaluating '
    'calibration dependence across haemodynamic monitoring devices.',
    space_after=Pt(12)
)

# Key contributions
add_para('Key Contributions', bold=True, space_after=Pt(4))

add_para(
    '\u2022 Analytical demonstration that zero calibration addresses only the location shift '
    '(u \u2192 0) component of the bias correction coefficient, not the scale shift (v)'
)

add_para(
    '\u2022 Simulation evidence that Bland\u2013Altman analysis fails to detect gain errors masked '
    'by offset cancellation, while CCC decomposition clearly identifies them'
)

add_para(
    '\u2022 Identification of a complete engineering pathway to calibration-free arterial '
    'pressure monitoring'
)

add_para(
    '\u2022 Proposal for reporting CCC with r/C\u2082 decomposition and pre-/post-calibration '
    'C\u2082 in all arterial pressure validation studies',
    space_after=Pt(12)
)

# Declarations
add_para('Declarations', bold=True, space_after=Pt(4))

add_para(
    '\u2022 This manuscript has not been published previously and is not under consideration '
    'elsewhere.'
)

add_para(
    '\u2022 All authors have approved the manuscript and agree with its submission to '
    'Physiological Measurement.'
)

add_para('\u2022 The authors declare no conflicts of interest.')

add_para('\u2022 [Funding statement to be completed by authors]', space_after=Pt(12))

# Closing
add_para(
    'We look forward to your consideration and welcome any feedback from reviewers.',
    space_after=Pt(18)
)

add_para('Sincerely,', space_after=Pt(18))

add_para('[Corresponding Author Name]')
add_para('[Affiliation]')
add_para('[Email]')
add_para('[ORCID]')

# Save
doc.save(OUTPATH)
print(f'Cover letter saved: {OUTPATH}')
